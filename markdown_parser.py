import json
import argparse
from docx import Document
from typing import Dict, List

def read_notebook(path: str) -> Dict:
    with open(path, encoding='utf8') as f:
        file = json.load(f)
    return file

def parse_notebook(file: Dict) -> Document:
    document = Document()
    for cell in file['cells']:
        if cell['cell_type'] == 'markdown':
            paragraph = parse_markdown_cell(cell['source'])
            paragraph = ' '.join([word for word in paragraph.split(' ') if len(word) < 100])
            document.add_paragraph(paragraph)
        elif cell['cell_type'] == 'code':
            code = ''.join(cell['source'])
            comments = extract_comments(code)
            if comments:
                comments = '$CODE COMMENTS\n' + '\n'.join(comments) + '\n$END CODE COMMENTS\n'
                document.add_paragraph(comments)
    return document

def extract_comments(code):
    comments = []
    in_comment_block = False
    lines = code.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            comments.append(line.lstrip('#').strip())
        if line.startswith('"""') or line.startswith("'''"):
            in_comment_block = not in_comment_block
            if in_comment_block:
                comments.append(line)
            else:
                comments[-1] += '\n' + line
        elif in_comment_block:
            comments[-1] += '\n' + line
    return comments

def parse_markdown_cell(lines: List) -> str:
    paragraph = ''
    for line in lines:
        if "![" not in line: # drop pictures
            paragraph += line
    return paragraph

def main():
    parser = argparse.ArgumentParser(description='Process command line arguments.')
    parser.add_argument('--path', type=str, help='notebook path for parsing')
    args = parser.parse_args()
    path = args.path

    file = read_notebook(path)
    document = parse_notebook(file)
    document.save('docx_file.docx')

if __name__ == "__main__":
    main()

